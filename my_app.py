from docx import Document 
from docx.shared import Inches


document = Document() 

# profile picture
document.add_picture('cv.png', width=Inches(1.0))


# name, phone number and email details
name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)


# about me
document.add_heading('About me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

document.save('cv.docx')
